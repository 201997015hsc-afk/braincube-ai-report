import streamlit as st
import pandas as pd
import plotly.express as px
import google.generativeai as genai
from docx import Document
import io
import datetime
import os # 파일 존재 확인용

# 디버깅용 코드 (잠시만 넣어보세요!)
import os
st.write("📍 현재 서버 위치의 파일들:", os.listdir())

# 1. 보안 및 페이지 설정
api_key = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=api_key)
st.set_page_config(page_title="브레인큐브 AI 리포트", layout="wide")

# 🔥 [수정 포인트] 로고 스마트 로더 (대소문자/파일위치 자동 찾기)
@st.cache_resource
def load_logo():
    # 깃허브에 올린 로고 파일 이름 후보들 (가장 흔한 오타/대소문자 조합)
    candidates = [
        'braincube_logo.png', 'Braincube_Logo.png', 'Logo.png', 'logo.png',
        'AI_Report/braincube_logo.png' # 폴더 안에 들었을 경우
    ]
    for cand in candidates:
        if os.path.exists(cand):
            return cand
    return None

# 🔥 [수정 포인트] 벤치마크 스마트 로더 (데이터 로드 및 CPA/CTR 계산)
@st.cache_data
def load_and_calculate_benchmark():
    candidates = ['benchmark.csv', 'Benchmark.csv', 'BENCHMARK.csv', 'AI_Report/benchmark.csv']
    target_file = None
    for cand in candidates:
        if os.path.exists(cand):
            target_file = cand
            break
    
    if target_file:
        try:
            b_df = pd.read_csv(target_file)
            # 벤치마크 엑셀 컬럼명 통일 (양끝 공백 제거 및 한글/영어 통일)
            b_df.columns = [c.strip() for c in b_df.columns]
            col_map_b = {
                '평균 CPA': ['평균 CPA', 'CPA', 'Avg CPA', '전환단가'],
                '평균 CTR': ['평균 CTR', 'CTR', 'Avg CTR', '클릭률']
            }
            for target, cands in col_map_b.items():
                for c in cands:
                    if c in b_df.columns:
                        b_df.rename(columns={c: target}, inplace=True)
                        break
            
            # 숫자 형식 변환
            b_df['평균 CPA'] = pd.to_numeric(b_df['평균 CPA'], errors='coerce')
            b_df['평균 CTR'] = pd.to_numeric(b_df['평균 CTR'], errors='coerce')
            
            # 전체 매체 평균값 산출
            return {
                'avg_cpa': b_df['평균 CPA'].mean(),
                'avg_ctr': b_df['평균 CTR'].mean()
            }
        except: return None
    return None

# 파일 로드 실행
logo_path = load_logo()
benchmark_averages = load_and_calculate_benchmark()

# 2. 로고 및 타이틀 (중앙 배치)
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    if logo_path:
        st.image(logo_path, width=300) 
    else:
        # 로고 못 찾았을 때 대안 텍스트 디자인
        st.markdown("<h1 style='text-align: center; color: orange;'>BRAINCUBE</h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center;'>AI 캠페인 분석 솔루션</h2>", unsafe_allow_html=True)
    st.write("---")

# 3. 데이터 입력 화면
if "uploaded_file" not in st.session_state:
    st.subheader("📂 데이터 분석 시작")
    uploaded_file = st.file_uploader("광고주 캠페인 로우데이터(CSV/XLSX)를 업로드하세요", type=["csv", "xlsx"])
else:
    uploaded_file = st.session_state.uploaded_file

if uploaded_file:
    # 데이터 읽기
    if uploaded_file.name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)

    # 엑셀 제목 스마트 매칭 로직 (이전 버전 유지)
    df.columns = [c.strip() for c in df.columns]
    col_map = {
        '소진액': ['소진액', '광고비', '비용', 'Spend', 'Cost'],
        'CTR': ['CTR', '클릭률'],
        'CPA': ['CPA', '전환단가'],
        '전환수': ['전환수', '전환', 'Conversions'],
        '매체': ['매체', '매체명', 'Media'],
        '소재명': ['소재명', '소재', 'Creative']
    }
    for target, candidates in col_map.items():
        for cand in candidates:
            if cand in df.columns:
                df.rename(columns={cand: target}, inplace=True)
                break

    # 필수 데이터 확인 및 전처리
    if '소진액' not in df.columns:
        st.error(f"⚠️ 파일에서 '소진액'이나 '광고비' 컬럼을 찾을 수 없습니다. (현재 컬럼명: {list(df.columns)})")
    else:
        # 데이터가 숫자 형식이 아닐 경우를 대비해 변환
        df['소진액'] = pd.to_numeric(df['소진액'], errors='coerce').fillna(0)
        df['전환수'] = pd.to_numeric(df.get('전환수', 0), errors='coerce').fillna(0)
        df['CPA'] = pd.to_numeric(df.get('CPA', 0), errors='coerce').fillna(0)
        df['CTR'] = pd.to_numeric(df.get('CTR', 0), errors='coerce').fillna(0)

        st.success(f"✅ '{uploaded_file.name}' 분석 완료!")

        # --- 대시보드 출력 ---
        tab1, tab2, tab3 = st.tabs(["📈 핵심 성과 & 벤치마크", "🎨 소재 효율 분석", "🤖 AI 전략 리포트"])

        with tab1:
            st.subheader("캠페인 종합 퍼포먼스 (vs 26년 평균)")
            m1_col, m2_col, m3_col, m4_col = st.columns(4)
            
            # 현재 캠페인 수치 계산
            current_ctr = df['CTR'].mean()
            current_cpa = df['CPA'].mean()
            
            # 🔥 [수정 포인트] 벤치마크 대비 Delta 계산 로직
            ctr_delta = None
            if benchmark_averages and 'avg_ctr' in benchmark_averages and benchmark_averages['avg_ctr'] > 0:
                ctr_diff = current_ctr - benchmark_averages['avg_ctr']
                # 상승/하락 비율 계산
                ctr_delta = f"{(ctr_diff / benchmark_averages['avg_ctr']) * 100:.1f}% (26년 평균 CTR 대비)"

            cpa_delta = None
            if benchmark_averages and 'avg_cpa' in benchmark_averages and benchmark_averages['avg_cpa'] > 0:
                cpa_diff = current_cpa - benchmark_averages['avg_cpa']
                # CPA는 낮을수록 좋으므로 inverse 처리 예정
                cpa_delta = f"{(cpa_diff / benchmark_averages['avg_cpa']) * 100:.1f}% (26년 평균 CPA 대비)"

            # 메트릭 표시 (벤치마크 데이터 있으면 Delta도 함께 표시)
            m1_col.metric("총 소진액", f"{df['소진액'].sum():,.0f}원")
            m2_col.metric("평균 CTR", f"{current_ctr:.2f}%", delta=ctr_delta)
            
            # CPA는 낮아지면(음수) 초록색, 높아지면(양수) 빨간색으로 표시 (delta_color="inverse")
            m3_col.metric("평균 CPA", f"{current_cpa:,.0f}원", delta=cpa_delta, delta_color="inverse")
            m4_col.metric("총 전환수", f"{df['전환수'].sum():,.0f}건")

            st.write("---")
            if '매체' in df.columns:
                st.subheader("매체별 예산 사용 및 CPA 현황")
                fig_media = px.bar(df, x='매체', y='소진액', color='CPA', title="매체별 소진액 (색상: CPA)", color_continuous_scale="Oranges")
                st.plotly_chart(fig_media, width='stretch')

        with tab2:
            if '소재명' in df.columns:
                st.subheader("CPA 기준 Top 5 위닝 소재")
                # CPA 0은 제외하고 소팅 (계산오류 방지)
                top_creatives = df[df['CPA'] > 0].sort_values(by='CPA', ascending=True).head(5)
                fig_creative = px.bar(top_creatives, x='소재명', y='CTR', color='CPA', title="Top 5 소재 CTR (색상: CPA 낮은 순)")
                st.plotly_chart(fig_creative, width='stretch')
                st.table(top_creatives[['소재명', 'CTR', 'CPA', '전환수']])
            else:
                st.info("소재별 데이터를 보려면 '소재명' 컬럼이 필요합니다.")

        with tab3:
            st.subheader("AI 다각도 캠페인 분석")
            if st.button("AI 전략 리포트 생성 시작"):
                with st.spinner("브레인큐브 AI가 데이터를 분석 중입니다..."):
                    
                    # AI에게 줄 데이터 요약 (벤치마크 정보 포함)
                    media_summary = df.groupby('매체').agg({'소진액':'sum', 'CPA':'mean', '전환수':'sum'}).to_string()
                    benchmark_info = ""
                    if benchmark_averages:
                        benchmark_info = f"\n(참고: 브레인큐브 26년 업종 평균 CPA: {benchmark_averages['avg_cpa']:.0f}원, 평균 CTR: {benchmark_averages['avg_ctr']:.2f}%)"
                    
                    model = genai.GenerativeModel('gemini-1.5-flash')
                    prompt = f"""
                    광고 대행사 브레인큐브의 시니어 마케터로서 아래 데이터를 기반으로 광고주 보고서를 작성해줘.
                    
                    1. 캠페인 현재 데이터:
                    {media_summary}
                    
                    2. 업종 벤치마크 데이터: {benchmark_info}
                    
                    요청하는 보고서 형식:
                    - 전월 성과 총평 (업종 평균과 비교하여 현재 캠페인의 해상도를 높여서 분석할 것)
                    - 고효율 매체 및 소재 발굴 (Data-driven 근거 제시)
                    - 차월 예산 재배분 및 소재 소구점 변경 전략 (Next Step)
                    """
                    
                    response = model.generate_content(prompt)
                    st.markdown(response.text)
                    
                    # 워드 다운로드 기능 (기존 유지)
                    doc = Document()
                    doc.add_heading('브레인큐브 AI 분석 리포트', 0)
                    doc.add_paragraph(response.text)
                    bio = io.BytesIO()
                    doc.save(bio)
                    st.download_button("📄 리포트 다운로드(.docx)", bio.getvalue(), f"Report_{datetime.date.today()}.docx")

else:
    # 깃허브에 파일이 없을 경우를 대비해 안내 텍스트 출력
    if not os.path.exists('benchmark.csv'):
        st.warning("⚠️ GitHub에 'benchmark.csv' 파일이 보이지 않습니다. 파일을 업로드하셔야 26년 실적 비교 기능이 활성화됩니다.")

    st.info("광고주 캠페인 리포트를 업로드하시면 AI 분석이 시작됩니다.")

st.markdown("---")
st.caption("© 2026 Braincube AI Marketing Solutions. All rights reserved.")
